from flask import Flask, request, render_template, redirect, url_for, flash, session, send_file
from flask_sqlalchemy import SQLAlchemy
import os
import numpy as np
import pandas as pd
import pickle
import torch
from PIL import Image
import torchvision.transforms.functional as TF
import CNN
import sklearn
import fitz
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_groq import ChatGroq
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
from flask_cors import CORS, cross_origin
import numpy as np
import pandas as pd
from datetime import datetime
import crops
import random
import firebase_admin
from firebase_admin import credentials, db


district_data = {
    "Anantapur": {
        "soil": "87.4% red soil, 12.5% black cotton soil, 0.1% problematic land",
        "major_crops": ["Groundnut", "Paddy", "Sweet Orange", "Mango", "Papaya", "Sunflower", "Green Gram", "Pigeon Pea", "Rice", "Millet"],
        "tips": ["Use protective irrigation for drought-prone areas.", "Practice multi-cropping to reduce risk."]
    },
    "Chittoor": {
        "soil": "70% dry land, 30% irrigated",
        "major_crops": ["Mango", "Cashew", "Tomato", "Papaya", "Sugarcane", "Peanut", "Rice", "Red Gram", "Sunflower"],
        "tips": ["Focus on mango and cashew cultivation.", "Use drip irrigation for water efficiency."]
    },
    "East Godavari": {
        "soil": "Alluvial soil",
        "major_crops": ["Paddy", "Sugarcane", "Coconut", "Mango", "Banana", "Cotton", "Green Gram", "Black Gram"],
        "tips": ["Paddy is the main crop; focus on irrigation management.", "Promote coconut and oil palm plantations."]
    },
    "Guntur": {
        "soil": "Fertile alluvial soil",
        "major_crops": ["Paddy", "Cotton", "Chilli", "Tobacco", "Maize", "Turmeric", "Banana", "Citrus", "Mango", "Guava"],
        "tips": ["Diversify crops with medicinal and aromatic plants.", "Improve storage for perishable crops like chillies."]
    },
    "Krishna": {
        "soil": "Fertile alluvial and coastal soil",
        "major_crops": ["Paddy", "Maize", "Cotton", "Sugarcane", "Peanuts", "Chillies", "Black Gram", "Mango", "Guava", "Cashew"],
        "tips": ["Utilize Krishna river water for efficient irrigation.", "Focus on vegetable cultivation for market profitability."]
    },
    "Kurnool": {
        "soil": "Red and black soils",
        "major_crops": ["Chickpea", "Groundnut", "Sunflower", "Rice", "Millet", "Cotton", "Pigeon Pea", "Black Gram", "Onion", "Tomato"],
        "tips": ["Improve post-harvest storage for onions and tomatoes.", "Encourage horticulture for better economic returns."]
    },
    "Prakasam": {
        "soil": "Sandy loam, red soil, black soil, sandy soil",
        "major_crops": ["Chickpea", "Paddy", "Red Gram", "Tobacco", "Cotton", "Sunflower", "Peanuts", "Orange", "Mango", "Sapota"],
        "tips": ["Diversify into horticulture and plantation crops.", "Use soil conservation techniques for sustainable farming."]
    },
    "Srikakulam": {
        "soil": "Coastal alluvial soil",
        "major_crops": ["Paddy", "Peanuts", "Sugarcane", "Green Gram", "Black Gram", "Horse Gram", "Mango", "Onion", "Chilli"],
        "tips": ["Encourage cashew and coconut plantations for higher income.", "Develop infrastructure for storage and processing of cash crops."]
    },
    "Nellore": {
        "soil": "Fertile alluvial soil",
        "major_crops": ["Paddy", "Black Gram", "Bengal Gram", "Groundnut", "Green Gram", "Red Gram"],
        "tips": ["Focus on paddy production with improved irrigation techniques.", "Promote pulses cultivation for crop diversification."]
    },
    "Visakhapatnam": {
        "soil": "Fertile and well-drained soil",
        "major_crops": ["Rice", "Ragi", "Bajra", "Jowar", "Groundnut", "Sugarcane", "Sesame", "Cotton", "Chillies", "Basmati Rice"],
        "tips": ["Experiment with high-value crops like strawberries and grapes.", "Utilize minor irrigation tanks efficiently."]
    },
    "Vizianagaram": {
        "soil": "Red and black soils",
        "major_crops": ["Rice", "Groundnut", "Mesta", "Sugarcane", "Cotton", "Maize", "Ragi", "Bajra", "Pulses"],
        "tips": ["Promote horticulture and post-harvest management.", "Encourage shade-net cultivation for vegetables."]
    },
    "West Godavari": {
        "soil": "Alluvial soil",
        "major_crops": ["Paddy", "Sugarcane", "Maize", "Tobacco", "Peanuts", "Pulses", "Sunflower", "Mango", "Coconut", "Turmeric"],
        "tips": ["Optimize paddy and sugarcane production with better irrigation.", "Promote vegetable and fruit cultivation for market value."]
    },
    "Kadapa": {
        "soil": "Red and black soils",
        "major_crops": ["Paddy", "Groundnut", "Red Gram", "Cotton", "Bengal Gram", "Mango", "Citrus", "Banana", "Papaya", "Turmeric"],
        "tips": ["Utilize riverbed areas for special crops like betel leaf.", "Encourage commercial cultivation of turmeric and onions."]
    }
}

# Initialize Flask app
app = Flask(__name__)
app.secret_key = 'your_secret_key'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)



cred = credentials.Certificate('serviceAccountKey.json')
firebase_admin.initialize_app(cred, {
    'databaseURL': 'https://smart-irrigation-63aab-default-rtdb.firebaseio.com/'
})


app.config['CORS_HEADERS'] = 'Content-Type'

cors = CORS(app, resources={r"/ticker": {"origins": "http://localhost:port"}})

commodity_dict = {
    "arhar": "static/Arhar.csv",
    "bajra": "static/Bajra.csv",
    "barley": "static/Barley.csv",
    "copra": "static/Copra.csv",
    "cotton": "static/Cotton.csv",
    "sesamum": "static/Sesamum.csv",
    "gram": "static/Gram.csv",
    "groundnut": "static/Groundnut.csv",
    "jowar": "static/Jowar.csv",
    "maize": "static/Maize.csv",
    "masoor": "static/Masoor.csv",
    "moong": "static/Moong.csv",
    "niger": "static/Niger.csv",
    "paddy": "static/Paddy.csv",
    "ragi": "static/Ragi.csv",
    "rape": "static/Rape.csv",
    "jute": "static/Jute.csv",
    "safflower": "static/Safflower.csv",
    "soyabean": "static/Soyabean.csv",
    "sugarcane": "static/Sugarcane.csv",
    "sunflower": "static/Sunflower.csv",
    "urad": "static/Urad.csv",
    "wheat": "static/Wheat.csv"
}

annual_rainfall = [29, 21, 37.5, 30.7, 52.6, 150, 299, 251.7, 179.2, 70.5, 39.8, 10.9]
base = {
    "Paddy": 1245.5,
    "Arhar": 3200,
    "Bajra": 1175,
    "Barley": 980,
    "Copra": 5100,
    "Cotton": 3600,
    "Sesamum": 4200,
    "Gram": 2800,
    "Groundnut": 3700,
    "Jowar": 1520,
    "Maize": 1175,
    "Masoor": 2800,
    "Moong": 3500,
    "Niger": 3500,
    "Ragi": 1500,
    "Rape": 2500,
    "Jute": 1675,
    "Safflower": 2500,
    "Soyabean": 2200,
    "Sugarcane": 2250,
    "Sunflower": 3700,
    "Urad": 4300,
    "Wheat": 1350

}
commodity_list = []


class Commodity:

    def __init__(self, csv_name):
        self.name = csv_name
        dataset = pd.read_csv(csv_name)
        self.X = dataset.iloc[:, :-1].values
        self.Y = dataset.iloc[:, 3].values

        #from sklearn.model_selection import train_test_split
        #X_train, X_test, Y_train, Y_test = train_test_split(X, Y, test_size=0.1, random_state=0)

        # Fitting decision tree regression to dataset
        from sklearn.tree import DecisionTreeRegressor
        depth = random.randrange(7,18)
        self.regressor = DecisionTreeRegressor(max_depth=depth)
        self.regressor.fit(self.X, self.Y)
        #y_pred_tree = self.regressor.predict(X_test)
        # fsa=np.array([float(1),2019,45]).reshape(1,3)
        # fask=regressor_tree.predict(fsa)

    def getPredictedValue(self, value):
        if value[1]>=2019:
            fsa = np.array(value).reshape(1, 3)
            #print(" ",self.regressor.predict(fsa)[0])
            return self.regressor.predict(fsa)[0]
        else:
            c=self.X[:,0:2]
            x=[]
            for i in c:
                x.append(i.tolist())
            fsa = [value[0], value[1]]
            ind = 0
            for i in range(0,len(x)):
                if x[i]==fsa:
                    ind=i
                    break
            #print(index, " ",ind)
            #print(x[ind])
            #print(self.Y[i])
            return self.Y[i]

    def getCropName(self):
        a = self.name.split('.')
        return a[0]

# Load models and data
crop_model = pickle.load(open('D:/Plant-Disease-Detection-main/Flask Deployed App/mod.pkl', 'rb'))
sc = pickle.load(open('D:/Plant-Disease-Detection-main/Flask Deployed App/sta.pkl', 'rb'))
ms = pickle.load(open('D:/Plant-Disease-Detection-main/Flask Deployed App/min.pkl', 'rb'))
disease_info = pd.read_csv('disease_info.csv', encoding='cp1252')
supplement_info = pd.read_csv('supplement_info.csv', encoding='cp1252')
disease_model = CNN.CNN(39)
disease_model.load_state_dict(torch.load("plant_disease_model_1_latest.pt", weights_only=True))
disease_model.eval()
dtr = pickle.load(open('dtr.pkl', 'rb'))
preprocessor = pickle.load(open('preprocessor.pkl', 'rb'))

# Define User model
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    email = db.Column(db.String(100), unique=True, nullable=False)
    password = db.Column(db.String(100), nullable=False)

# Create tables
with app.app_context():
    db.create_all()

# Global variables for QA chain
qa = None
texts = []
chat_history = []

# Function to extract text from a PDF file
def load_pdf_text(pdf_path):
    doc = fitz.open(pdf_path)
    text_data = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        text_data.append(page.get_text("text"))
    return text_data

# Function to fetch and extract text from a URL
def fetch_url_content(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            if url.lower().endswith(".pdf"):
                with open("temp.pdf", "wb") as f:
                    f.write(response.content)
                return load_pdf_text("temp.pdf")
            else:
                soup = BeautifulSoup(response.text, "html.parser")
                return [soup.get_text()]
        else:
            return []
    except Exception as e:
        return []

# Function to initialize the QA chain
def initialize_qa_chain(texts):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    text_chunks = text_splitter.create_documents(texts)
    text_embeddings = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
    text_db = FAISS.from_documents(text_chunks, text_embeddings)
    os.environ['GROQ_API_KEY'] = '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'
    llm = ChatGroq(model_name='llama-3.1-8b-instant')
    memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
    retriever = text_db.as_retriever(search_type="similarity", search_kwargs={"k": 3})
    qa = ConversationalRetrievalChain.from_llm(llm=llm, memory=memory, retriever=retriever)
    return qa

# Function to create a boxed table in a Word document
def create_boxed_table(doc, label, value):
    table = doc.add_table(rows=1, cols=len(value))
    table.autofit = False
    table.allow_autofit = False
    for i, char in enumerate(value):
        cell = table.cell(0, i)
        cell.text = char
        cell.width = Pt(20)
        cell.paragraphs[0].alignment = 1
        cell.vertical_alignment = 1
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    doc.add_paragraph(label, style='Intense Quote')

# Function to generate an agricultural loan application document
def generate_agri_loan_application(personal_info, land_details, loan_details, bank_details, photo, signature):
    doc = Document()
    doc.add_heading('Rythu-Mitra Agricultural Loan Application for Farmers', 0)
    if photo is not None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(photo, width=Inches(1.5))
        paragraph.alignment = 2
    doc.add_heading('Personal Information', level=1)
    create_boxed_table(doc, "Full Name:", personal_info['name'])
    create_boxed_table(doc, "Date of Birth:", str(personal_info['dob']))
    create_boxed_table(doc, "Identity Proof Type:", personal_info['identity_proof_type'])
    create_boxed_table(doc, "Identity Proof Number:", personal_info['identity_proof'])
    create_boxed_table(doc, "Address Proof:", personal_info['address_proof'])
    create_boxed_table(doc, "Mobile Number:", personal_info['mobile'])
    create_boxed_table(doc, "Email ID:", personal_info['email'])
    doc.add_heading('Land & Farming Details', level=1)
    create_boxed_table(doc, "Land Area:", land_details['land_area'])
    create_boxed_table(doc, "Location:", land_details['location'])
    create_boxed_table(doc, "Crop Type:", land_details['crop_type'])
    create_boxed_table(doc, "Past Yield:", land_details['past_yield'])
    create_boxed_table(doc, "Water Source:", land_details['water_source'])
    doc.add_heading('Loan-Related Details', level=1)
    create_boxed_table(doc, "Loan Amount Required:", loan_details['loan_amount'])
    create_boxed_table(doc, "Purpose of Loan:", loan_details['purpose'])
    create_boxed_table(doc, "Repayment Period:", loan_details['repayment_period'])
    create_boxed_table(doc, "Previous Loan Details:", loan_details['previous_loan'])
    doc.add_heading('Bank Details', level=1)
    create_boxed_table(doc, "Bank Account Number:", bank_details['account_number'])
    doc.add_heading('Declaration', level=1)
    doc.add_paragraph("I certify that the information given above and in the enclosures are true in all respects and that this shall form the basis of loan application.")
    if signature is not None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(signature, width=Inches(1.5))
        paragraph.alignment = 2
    doc.add_paragraph(f"Name and Signature/Thumb Impression of the Applicant: {personal_info['name']}")
    doc.add_paragraph(f"Place: {personal_info['city']}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Home Page
@app.route('/')
def home_page():
    return render_template('home.html')

@app.route('/forecast')
def hott():
    return render_template('forecast.html')

@app.route('/forecast')
def ind():
    # Fetch data from Firebase
    ref = db.reference()
    sensor_data = ref.get()
    
    # Extract sensor data
    temperature = sensor_data.get('temperature', 'N/A')  # Use 'N/A' as default if key doesn't exist
    humidity = sensor_data.get('humidity', 'N/A')
    soil_moisture1 = sensor_data.get('moisture1', 'N/A')
    soil_moisture2 = sensor_data.get('moisture2', 'N/A')
    soil_moisture3 = sensor_data.get('moisture3', 'N/A')
    
    # Pass data to template
    return render_template('forecast.html', 
                           temperature=temperature, 
                           humidity=humidity, 
                           soil_moisture1=soil_moisture1, 
                           soil_moisture2=soil_moisture2, 
                           soil_moisture3=soil_moisture3)


@app.route('/district')
def ho():
    return render_template('district.html')

@app.route("/reco", methods=["POST"])
def rec():
    district = request.form.get("district")
    data = district_data.get(district, {})
    return render_template("discrop.html", district=district, data=data)

# Signup Route
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        name = request.form['name']
        email = request.form['email']
        password = request.form['password']
        confirm_password = request.form['confirm_password']
        if password != confirm_password:
            flash('Passwords do not match!')
            return redirect(url_for('signup'))
        new_user = User(name=name, email=email, password=password)
        db.session.add(new_user)
        db.session.commit()
        flash('Account created successfully! Please login.')
        return redirect(url_for('login'))
    return render_template('signup.html')

# Login Route
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        user = User.query.filter_by(email=email, password=password).first()
        if user:
            session['user_id'] = user.id
            return render_template('home.html')
        else:
            flash('Invalid email or password')
            return render_template('login.html')
    return render_template('login.html')



@app.route('/new')
def index():
    context = {
        "top5": TopFiveWinners(),
        "bottom5": TopFiveLosers(),
        "sixmonths": SixMonthsForecast()
    }
    return render_template('new.html', context=context)


@app.route('/commodity/<name>')
def crop_profile(name):
    max_crop, min_crop, forecast_crop_values = TwelveMonthsForecast(name)
    prev_crop_values = TwelveMonthPrevious(name)
    forecast_x = [i[0] for i in forecast_crop_values]
    forecast_y = [i[1] for i in forecast_crop_values]
    previous_x = [i[0] for i in prev_crop_values]
    previous_y = [i[1] for i in prev_crop_values]
    current_price = CurrentMonth(name)
    #print(max_crop)
    #print(min_crop)
    #print(forecast_crop_values)
    #print(prev_crop_values)
    #print(str(forecast_x))
    crop_data = crops.crop(name)
    context = {
        "name":name,
        "max_crop": max_crop,
        "min_crop": min_crop,
        "forecast_values": forecast_crop_values,
        "forecast_x": str(forecast_x),
        "forecast_y":forecast_y,
        "previous_values": prev_crop_values,
        "previous_x":previous_x,
        "previous_y":previous_y,
        "current_price": current_price,
        "image_url":crop_data[0],
        "prime_loc":crop_data[1],
        "type_c":crop_data[2],
        "export":crop_data[3]
    }
    return render_template('commodity.html', context=context)

@app.route('/ticker/<item>/<number>')
@cross_origin(origin='localhost',headers=['Content- Type','Authorization'])
def ticker(item, number):
    n = int(number)
    i = int(item)
    data = SixMonthsForecast()
    context = str(data[n][i])

    if i == 2 or i == 5:
        context = '₹' + context
    elif i == 3 or i == 6:

        context = context + '%'

    #print('context: ', context)
    return context


def TopFiveWinners():
    current_month = datetime.now().month
    current_year = datetime.now().year
    current_rainfall = annual_rainfall[current_month - 1]
    prev_month = current_month - 1
    prev_rainfall = annual_rainfall[prev_month - 1]
    current_month_prediction = []
    prev_month_prediction = []
    change = []

    for i in commodity_list:
        current_predict = i.getPredictedValue([float(current_month), current_year, current_rainfall])
        current_month_prediction.append(current_predict)
        prev_predict = i.getPredictedValue([float(prev_month), current_year, prev_rainfall])
        prev_month_prediction.append(prev_predict)
        change.append((((current_predict - prev_predict) * 100 / prev_predict), commodity_list.index(i)))
    sorted_change = change
    sorted_change.sort(reverse=True)
    # print(sorted_change)
    to_send = []
    for j in range(0, 5):
        perc, i = sorted_change[j]
        name = commodity_list[i].getCropName().split('/')[1]
        to_send.append([name, round((current_month_prediction[i] * base[name]) / 100, 2), round(perc, 2)])
    #print(to_send)
    return to_send


def TopFiveLosers():
    current_month = datetime.now().month
    current_year = datetime.now().year
    current_rainfall = annual_rainfall[current_month - 1]
    prev_month = current_month - 1
    prev_rainfall = annual_rainfall[prev_month - 1]
    current_month_prediction = []
    prev_month_prediction = []
    change = []

    for i in commodity_list:
        current_predict = i.getPredictedValue([float(current_month), current_year, current_rainfall])
        current_month_prediction.append(current_predict)
        prev_predict = i.getPredictedValue([float(prev_month), current_year, prev_rainfall])
        prev_month_prediction.append(prev_predict)
        change.append((((current_predict - prev_predict) * 100 / prev_predict), commodity_list.index(i)))
    sorted_change = change
    sorted_change.sort()
    to_send = []
    for j in range(0, 5):
        perc, i = sorted_change[j]
        name = commodity_list[i].getCropName().split('/')[1]
        to_send.append([name, round((current_month_prediction[i] * base[name]) / 100, 2), round(perc, 2)])
   # print(to_send)
    return to_send



def SixMonthsForecast():
    month1=[]
    month2=[]
    month3=[]
    month4=[]
    month5=[]
    month6=[]
    for i in commodity_list:
        crop=SixMonthsForecastHelper(i.getCropName())
        k=0
        for j in crop:
            time = j[0]
            price = j[1]
            change = j[2]
            if k==0:
                month1.append((price,change,i.getCropName().split("/")[1],time))
            elif k==1:
                month2.append((price,change,i.getCropName().split("/")[1],time))
            elif k==2:
                month3.append((price,change,i.getCropName().split("/")[1],time))
            elif k==3:
                month4.append((price,change,i.getCropName().split("/")[1],time))
            elif k==4:
                month5.append((price,change,i.getCropName().split("/")[1],time))
            elif k==5:
                month6.append((price,change,i.getCropName().split("/")[1],time))
            k+=1
    month1.sort()
    month2.sort()
    month3.sort()
    month4.sort()
    month5.sort()
    month6.sort()
    crop_month_wise=[]
    crop_month_wise.append([month1[0][3],month1[len(month1)-1][2],month1[len(month1)-1][0],month1[len(month1)-1][1],month1[0][2],month1[0][0],month1[0][1]])
    crop_month_wise.append([month2[0][3],month2[len(month2)-1][2],month2[len(month2)-1][0],month2[len(month2)-1][1],month2[0][2],month2[0][0],month2[0][1]])
    crop_month_wise.append([month3[0][3],month3[len(month3)-1][2],month3[len(month3)-1][0],month3[len(month3)-1][1],month3[0][2],month3[0][0],month3[0][1]])
    crop_month_wise.append([month4[0][3],month4[len(month4)-1][2],month4[len(month4)-1][0],month4[len(month4)-1][1],month4[0][2],month4[0][0],month4[0][1]])
    crop_month_wise.append([month5[0][3],month5[len(month5)-1][2],month5[len(month5)-1][0],month5[len(month5)-1][1],month5[0][2],month5[0][0],month5[0][1]])
    crop_month_wise.append([month6[0][3],month6[len(month6)-1][2],month6[len(month6)-1][0],month6[len(month6)-1][1],month6[0][2],month6[0][0],month6[0][1]])

   # print(crop_month_wise)
    return crop_month_wise

def SixMonthsForecastHelper(name):
    current_month = datetime.now().month
    current_year = datetime.now().year
    current_rainfall = annual_rainfall[current_month - 1]
    name = name.split("/")[1]
    name = name.lower()
    commodity = commodity_list[0]
    for i in commodity_list:
        if name == str(i):
            commodity = i
            break
    month_with_year = []
    for i in range(1, 7):
        if current_month + i <= 12:
            month_with_year.append((current_month + i, current_year, annual_rainfall[current_month + i - 1]))
        else:
            month_with_year.append((current_month + i - 12, current_year + 1, annual_rainfall[current_month + i - 13]))
    wpis = []
    current_wpi = commodity.getPredictedValue([float(current_month), current_year, current_rainfall])
    change = []

    for m, y, r in month_with_year:
        current_predict = commodity.getPredictedValue([float(m), y, r])
        wpis.append(current_predict)
        change.append(((current_predict - current_wpi) * 100) / current_wpi)

    crop_price = []
    for i in range(0, len(wpis)):
        m, y, r = month_with_year[i]
        x = datetime(y, m, 1)
        x = x.strftime("%b %y")
        crop_price.append([x, round((wpis[i]* base[name.capitalize()]) / 100, 2) , round(change[i], 2)])

   # print("Crop_Price: ", crop_price)
    return crop_price

def CurrentMonth(name):
    current_month = datetime.now().month
    current_year = datetime.now().year
    current_rainfall = annual_rainfall[current_month - 1]
    name = name.lower()
    commodity = commodity_list[0]
    for i in commodity_list:
        if name == str(i):
            commodity = i
            break
    current_wpi = commodity.getPredictedValue([float(current_month), current_year, current_rainfall])
    current_price = (base[name.capitalize()]*current_wpi)/100
    return current_price

def TwelveMonthsForecast(name):
    current_month = datetime.now().month
    current_year = datetime.now().year
    current_rainfall = annual_rainfall[current_month - 1]
    name = name.lower()
    commodity = commodity_list[0]
    for i in commodity_list:
        if name == str(i):
            commodity = i
            break
    month_with_year = []
    for i in range(1, 13):
        if current_month + i <= 12:
            month_with_year.append((current_month + i, current_year, annual_rainfall[current_month + i - 1]))
        else:
            month_with_year.append((current_month + i - 12, current_year + 1, annual_rainfall[current_month + i - 13]))
    max_index = 0
    min_index = 0
    max_value = 0
    min_value = 9999
    wpis = []
    current_wpi = commodity.getPredictedValue([float(current_month), current_year, current_rainfall])
    change = []

    for m, y, r in month_with_year:
        current_predict = commodity.getPredictedValue([float(m), y, r])
        if current_predict > max_value:
            max_value = current_predict
            max_index = month_with_year.index((m, y, r))
        if current_predict < min_value:
            min_value = current_predict
            min_index = month_with_year.index((m, y, r))
        wpis.append(current_predict)
        change.append(((current_predict - current_wpi) * 100) / current_wpi)

    max_month, max_year, r1 = month_with_year[max_index]
    min_month, min_year, r2 = month_with_year[min_index]
    min_value = min_value * base[name.capitalize()] / 100
    max_value = max_value * base[name.capitalize()] / 100
    crop_price = []
    for i in range(0, len(wpis)):
        m, y, r = month_with_year[i]
        x = datetime(y, m, 1)
        x = x.strftime("%b %y")
        crop_price.append([x, round((wpis[i]* base[name.capitalize()]) / 100, 2) , round(change[i], 2)])
   # print("forecasr", wpis)
    x = datetime(max_year,max_month,1)
    x = x.strftime("%b %y")
    max_crop = [x, round(max_value,2)]
    x = datetime(min_year, min_month, 1)
    x = x.strftime("%b %y")
    min_crop = [x, round(min_value,2)]

    return max_crop, min_crop, crop_price


def TwelveMonthPrevious(name):
    name = name.lower()
    current_month = datetime.now().month
    current_year = datetime.now().year
    current_rainfall = annual_rainfall[current_month - 1]
    commodity = commodity_list[0]
    wpis = []
    crop_price = []
    for i in commodity_list:
        if name == str(i):
            commodity = i
            break
    month_with_year = []
    for i in range(1, 13):
        if current_month - i >= 1:
            month_with_year.append((current_month - i, current_year, annual_rainfall[current_month - i - 1]))
        else:
            month_with_year.append((current_month - i + 12, current_year - 1, annual_rainfall[current_month - i + 11]))

    for m, y, r in month_with_year:
        current_predict = commodity.getPredictedValue([float(m), 2013, r])
        wpis.append(current_predict)

    for i in range(0, len(wpis)):
        m, y, r = month_with_year[i]
        x = datetime(y,m,1)
        x = x.strftime("%b %y")
        crop_price.append([x, round((wpis[i]* base[name.capitalize()]) / 100, 2)])
   # print("previous ", wpis)
    new_crop_price =[]
    for i in range(len(crop_price)-1,-1,-1):
        new_crop_price.append(crop_price[i])
    return new_crop_price

# Logout Route
@app.route('/logout')
def logout():
    session.pop('user_id', None)
    return redirect(url_for('home_page'))

# Contact Page
@app.route('/contact')
def contact():
    return render_template('contact-us.html')

# Crop Recommendation
@app.route('/recomendation')
def recomendation():
    return render_template('recomendation1.html')

# AI Engine Page
@app.route('/index')
def ai_engine_page():
    return render_template('index.html')

# Mobile Device Detected Page
@app.route('/mobile-device')
def mobile_device_detected_page():
    return render_template('mobile-device.html')

# Crop Prediction
@app.route('/predict', methods=['POST'])
def predict():
    N = request.form['Nitrogen']
    P = request.form['Phosporus']
    K = request.form['Potassium']
    temp = request.form['Temperature']
    humidity = request.form['Humidity']
    ph = request.form['Ph']
    rainfall = request.form['Rainfall']
    feature_list = [N, P, K, temp, humidity, ph, rainfall]
    single_pred = np.array(feature_list).reshape(1, -1)
    scaled_features = ms.transform(single_pred)
    final_features = sc.transform(scaled_features)
    prediction = crop_model.predict(final_features)
    crop_dict = {1: "Rice", 2: "Maize", 3: "Jute", 4: "Cotton", 5: "Coconut", 6: "Papaya", 7: "Orange",
                 8: "Apple", 9: "Muskmelon", 10: "Watermelon", 11: "Grapes", 12: "Mango", 13: "Banana",
                 14: "Pomegranate", 15: "Lentil", 16: "Blackgram", 17: "Mungbean", 18: "Mothbeans",
                 19: "Pigeonpeas", 20: "Kidneybeans", 21: "Chickpea", 22: "Coffee"}
    if prediction[0] in crop_dict:
        crop = crop_dict[prediction[0]]
        result = "{} is the best crop to be cultivated right there".format(crop)
    else:
        result = "Sorry, we could not determine the best crop to be cultivated with the provided data."
    return render_template('recomendation1.html', result=result)

# Image Prediction for Disease
def prediction(image_path):
    image = Image.open(image_path)
    image = image.resize((224, 224))
    input_data = TF.to_tensor(image)
    input_data = input_data.view((-1, 3, 224, 224))
    output = disease_model(input_data)
    output = output.detach().numpy()
    index = np.argmax(output)
    return index

# Submit Image for Disease Detection
@app.route('/submit', methods=['GET', 'POST'])
def submit():
    if request.method == 'POST':
        image = request.files['image']
        filename = image.filename
        file_path = os.path.join('static/uploads', filename)
        image.save(file_path)
        pred = prediction(file_path)
        title = disease_info['disease_name'][pred]
        description = disease_info['description'][pred]
        prevent = disease_info['Possible Steps'][pred]
        image_url = disease_info['image_url'][pred]
        supplement_name = supplement_info['supplement name'][pred]
        supplement_image_url = supplement_info['supplement image'][pred]
        supplement_buy_link = supplement_info['buy link'][pred]
        return render_template('submit.html', title=title, desc=description, prevent=prevent,
                               image_url=image_url, pred=pred, sname=supplement_name, simage=supplement_image_url, buy_link=supplement_buy_link)

# Market Page
@app.route('/market', methods=['GET', 'POST'])
def market():
    return render_template('market.html', supplement_image=list(supplement_info['supplement image']),
                           supplement_name=list(supplement_info['supplement name']), disease=list(disease_info['disease_name']), buy=list(supplement_info['buy link']))

# Yield Prediction Page
@app.route('/yield')
def yield_prediction_page():
    return render_template('yield.html')

@app.route('/bot')
def agribot_is():
    return render_template('agribot.html')   

@app.route('/yiee')
def yield_prediction():
    return render_template('yiee.html')

# Predict Yield based on form inputs
@app.route("/vasee", methods=['POST'])
def vasee():
    if request.method == 'POST':
        Year = request.form['Year']
        average_rain_fall_mm_per_year = request.form['average_rain_fall_mm_per_year']
        pesticides_tonnes = request.form['pesticides_tonnes']
        avg_temp = request.form['avg_temp']
        Area = request.form['Area']
        Item = request.form['Item']
        features = np.array([[Year, average_rain_fall_mm_per_year, pesticides_tonnes, avg_temp, Area, Item]], dtype=object)
        transformed_features = preprocessor.transform(features)
        prediction = dtr.predict(transformed_features).reshape(1, -1)
        return render_template('yiee.html', prediction=prediction[0][0])

# Loan Application and QA Routes
@app.route('/loan', methods=['GET', 'POST'])
def loan():
    global qa, texts
    if request.method == 'POST':
        option = request.form.get('option')
        if option == "Upload PDF":
            uploaded_file = request.files['file']
            if uploaded_file:
                with open("uploaded_file.pdf", "wb") as f:
                    f.write(uploaded_file.read())
                texts = load_pdf_text("uploaded_file.pdf")
                qa = initialize_qa_chain(texts)
                return redirect(url_for('extracted_text'))
        elif option == "Submit URL":
            url = request.form.get('url')
            if url:
                texts = fetch_url_content(url)
                if texts:
                    qa = initialize_qa_chain(texts)
                    return redirect(url_for('extracted_text'))
        elif option == "Agri Loan Application":
            return redirect(url_for('agri_loan_application'))
    return render_template('loan.html', option=None)

@app.route('/extracted_text', methods=['GET', 'POST'])
def extracted_text():
    global qa, texts
    selected_page = request.args.get('page', default=1, type=int)
    return render_template('extracted_text.html', texts=texts, selected_page=selected_page)

@app.route('/ask', methods=['POST'])
def ask_question():
    global qa, texts, chat_history
    query = request.form.get('query')
    if query and qa:
        chat_history.append({"role": "Human", "content": query})
        result = qa({"question": query, "chat_history": chat_history})
        answer = result['answer']
        chat_history.append({"role": "AI", "content": answer})
        answer_points = answer.split(". ")
        return render_template('extracted_text.html', texts=texts, answer_points=answer_points, selected_page=1)
    return redirect(url_for('extracted_text'))

@app.route('/agri_loan_application', methods=['GET', 'POST'])
def agri_loan_application():
    if request.method == 'POST':
        personal_info = {
            "name": request.form.get('name'),
            "dob": request.form.get('dob'),
            "identity_proof_type": request.form.get('identity_proof_type'),
            "identity_proof": request.form.get('identity_proof'),
            "address_proof": request.form.get('address_proof'),
            "mobile": request.form.get('mobile'),
            "email": request.form.get('email'),
            "city": request.form.get('location').split(",")[-1].strip() if request.form.get('location') else ""
        }
        land_details = {
            "land_area": request.form.get('land_area'),
            "location": request.form.get('location'),
            "crop_type": request.form.get('crop_type'),
            "past_yield": request.form.get('past_yield'),
            "water_source": request.form.get('water_source')
        }
        loan_details = {
            "loan_amount": request.form.get('loan_amount'),
            "purpose": request.form.get('purpose'),
            "repayment_period": request.form.get('repayment_period'),
            "previous_loan": request.form.get('previous_loan')
        }
        bank_details = {
            "account_number": request.form.get('account_number')
        }
        photo = request.files['photo']
        signature = request.files['signature']
        doc_bytes = generate_agri_loan_application(personal_info, land_details, loan_details, bank_details, photo, signature)
        return send_file(
            doc_bytes,
            as_attachment=True,
            download_name="agri_loan_application.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    return render_template('agri_loan_application.html')

if __name__ == '__main__':
     arhar = Commodity(commodity_dict["arhar"])
     commodity_list.append(arhar)
     bajra = Commodity(commodity_dict["bajra"])
     commodity_list.append(bajra)
     barley = Commodity(commodity_dict["barley"])
     commodity_list.append(barley)
     copra = Commodity(commodity_dict["copra"])
     commodity_list.append(copra)
     cotton = Commodity(commodity_dict["cotton"])
     commodity_list.append(cotton)
     sesamum = Commodity(commodity_dict["sesamum"])
     commodity_list.append(sesamum)
     gram = Commodity(commodity_dict["gram"])
     commodity_list.append(gram)
     groundnut = Commodity(commodity_dict["groundnut"])
     commodity_list.append(groundnut)
     jowar = Commodity(commodity_dict["jowar"])
     commodity_list.append(jowar)
     maize = Commodity(commodity_dict["maize"])
     commodity_list.append(maize)
     masoor = Commodity(commodity_dict["masoor"])
     commodity_list.append(masoor)
     moong = Commodity(commodity_dict["moong"])
     commodity_list.append(moong)
     niger = Commodity(commodity_dict["niger"])
     commodity_list.append(niger)
     paddy = Commodity(commodity_dict["paddy"])
     commodity_list.append(paddy)
     ragi = Commodity(commodity_dict["ragi"])
     commodity_list.append(ragi)
     rape = Commodity(commodity_dict["rape"])
     commodity_list.append(rape)
     jute = Commodity(commodity_dict["jute"])
     commodity_list.append(jute)
     safflower = Commodity(commodity_dict["safflower"])
     commodity_list.append(safflower)
     soyabean = Commodity(commodity_dict["soyabean"])
     commodity_list.append(soyabean)
     sugarcane = Commodity(commodity_dict["sugarcane"])
     commodity_list.append(sugarcane)
     sunflower = Commodity(commodity_dict["sunflower"])
     commodity_list.append(sunflower)
     urad = Commodity(commodity_dict["urad"])
     commodity_list.append(urad)
     wheat = Commodity(commodity_dict["wheat"])
     commodity_list.append(wheat)

     app.run()
